import os
import docx
import re


def count_words_in_folder(folder_path):
    total_words = 0
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.endswith(('.docx', '.doc')):
                file_path = os.path.join(root, file)
                doc = docx.Document(file_path)
                for paragraph in doc.paragraphs:
                    chinese_characters = re.findall(r'[\u4e00-\u9fa5]', paragraph.text)
                    english_words = re.findall(r'\b\w+\b', paragraph.text)
                    total_words += len(chinese_characters) + len(english_words)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            chinese_characters = re.findall(r'[\u4e00-\u9fa5]', cell.text)
                            english_words = re.findall(r'\b\w+\b', cell.text)
                            total_words += len(chinese_characters) + len(english_words)
    return total_words

# 修改文件夹的路径：按本脚本的路径写入相对路径，比如说："./data/word_count"，即该脚本路径的data文件夹中的word_count文件夹中
folder_path = "./data/word_count/"
dir_list = os.listdir(folder_path)

for dir in dir_list:
    word_count = count_words_in_folder(folder_path+dir)
    print(f"{dir}文件夹中所有docx和doc格式文档中的字数为：{word_count}")
