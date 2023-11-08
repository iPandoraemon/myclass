import pandas as pd
import os
import re
from docx import Document
from docxcompose.composer import Composer


def read_reports(folder_path: str, student_list_file: str,
                 class_id: int) -> pd.DataFrame:
    exp_list = []
    for root, dirs, files in os.walk(folder_path):
        for dir_name in dirs:
            exp_list.append(dir_name)

    reports = pd.read_excel(student_list_file, usecols=['学号', '姓名'])
    student_num = reports.shape[0]
    reports['学号'] = reports['学号'].astype(str)
    # for exp in exp_list:
    #     reports[exp] = [0 for i in range(student_num)]

    for root, dirs, files in os.walk(folder_path):
        students = {}
        for file in files:
            split_root = re.split(r"[\\/]", root)
            # root.split('\\｜//|/|\')
            exp_name = split_root[-1]
            split_file = re.split('[-.(]', file)
            if split_file[-1] != 'docx':
                continue
            if is_number(split_file[0]):
                if not is_number(split_file[1]):
                    student_name = split_file[1]
                    student_id = split_file[0]
            elif re.findall(r'\d+', split_file[0]):
                student_name = re.findall(r'\D+', split_file[0])[0]
                student_id = re.findall(r'\d+', split_file[0])[0]
            else:
                student_name = split_file[0]
                student_id = split_file[1]
            student_name = re.findall(r'\D+', student_name)[0]
            student_name = student_name.strip()
            student_id = re.findall(r'\d+', student_id)[0]
            student_id = student_id.strip()
            if len(student_id) < 10:
                student_id = '2000502' + class_id + student_id[-2:]
                print(root+file+':'+student_id)
            students.setdefault(student_id, [student_id, 1])
            # reports.loc[reports['id'] == student_id, ['name']] = student_name
            # reports.loc[reports['id'] == student_id, [exp_name]] = 1
        if students:
            df = pd.DataFrame(students)
            df = df.T
            df = df.rename(columns={0: '学号', 1: exp_name})
            # df = df.rename(columns={0: '学号', 1: '姓名', 2: exp_name})
            reports = reports.merge(df, on='学号', how='left')
            print(exp_name + ' 结束')
    return reports


def output_list(reports: pd.DataFrame, output_file: str, output_sheet: str):
    with pd.ExcelWriter(output_file) as writer:
        reports.to_excel(writer, sheet_name=output_sheet)


def is_number(string):
    return bool(re.match('^[0-9]+$', string))


def is_number_and_other(string):
    return all(char.isdigit() or not char.isalnum() for char in string)


def merge_reports(original_docx_path: str, new_docx_path: str):
    all_file_path = []
    for file_name in os.listdir(original_docx_path):
        all_file_path.append(f'{original_docx_path}/{file_name}')

    first_document = Document(all_file_path[0])
    first_document.add_page_break()
    middle_new_docx = Composer(first_document)

    for index, word in enumerate(all_file_path[1:]):
        word_document = Document(word)
        if index != len(all_file_path) - 2:
            word_document.add_page_break()
            middle_new_docx.append(word_document)

    middle_new_docx.save(new_docx_path)


def combine_word_documents(file_paths, output_path):
    # Create a new document object
    merged_document = Document()

    # Loop through each document and add the content to the new document
    for file_path in file_paths:
        sub_doc = Document(file_path)

        # Add the content of the sub document to the merged document
        for element in sub_doc.element.body:
            merged_document.element.body.append(element)

        # Add a page break to signal the start of a new document
        merged_document.add_page_break()

        # Add the file name to the first paragraph of the new page
        last_paragraph = merged_document.paragraphs[-1]
        file_name = file_path.split('/')[-1]
        last_paragraph.add_run(file_name)

    # Add the file name to the first paragraph of the merged document
    first_paragraph = merged_document.paragraphs[0]
    first_paragraph.add_run('Merged Document')

    # Save the merged document
    merged_document.save(output_path)


def get_file_names(folder_path):
    # Get a list of all the files in the folder
    file_names = os.listdir(folder_path)

    # Return the list of file names
    return file_names
